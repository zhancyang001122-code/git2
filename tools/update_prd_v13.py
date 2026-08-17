from pathlib import Path
import shutil

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "ArchFlow_AI_PRD_v1.2.docx"
OUTPUT = ROOT / "docs" / "ArchFlow_AI_PRD_v1.3.docx"


def replace_runs(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_cell(cell, text):
    replace_runs(cell.paragraphs[0], text)
    for paragraph in cell.paragraphs[1:]:
        replace_runs(paragraph, "")


def keep_table_rows_intact(table):
    for index, row in enumerate(table.rows):
        properties = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "true")
        properties.append(cant_split)
        if index == 0:
            repeat_header = OxmlElement("w:tblHeader")
            repeat_header.set(qn("w:val"), "true")
            properties.append(repeat_header)


def update_document():
    shutil.copyfile(SOURCE, OUTPUT)
    document = Document(OUTPUT)
    paragraphs = document.paragraphs

    replace_runs(
        paragraphs[2],
        "版本：V1.3  |  产品形态：企业设计师 Web 工作台  |  文档状态：7 天面试演示 POC",
    )
    replace_runs(
        paragraphs[112],
        "当前目标是在约 7 天内完成可用于面试展示的 Web MVP。方案灵感、方案设计与 AI 渲染接入 ArchFlow 预设的两类真实 API；访问者界面只填写语言大模型与生图大模型两把 API Key。默认预设为 gpt-5-mini 与 gpt-image-2，密钥只保存在 sessionStorage，关闭标签页后清除。图纸美化、AI 建模与 AI 汇报保持高完成度演示数据。",
    )
    replace_runs(
        paragraphs[118],
        "能清楚说明两类内置模型预设、三条真实生成链路、三条演示链路，以及静态公开演示与服务端免填 Key 方案的边界。",
    )
    replace_runs(paragraphs[16], "我的资产")
    replace_runs(
        paragraphs[17],
        "整体流程：方案灵感 → 方案设计 → 图纸美化 / AI 建模 / AI 渲染 → AI 汇报 → “我的资产”沉淀历史成果。各功能可独立使用，不要求强制串联。",
    )

    replace_cell(document.tables[3].rows[4].cells[0], "脱敏演示身份与两把 API Key 配置")

    for table in document.tables:
        keep_table_rows_intact(table)

    for section in document.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                if paragraph.text.startswith("ArchFlow AI · PRD"):
                    replace_runs(paragraph, "ArchFlow AI · PRD V1.3")
                elif "V1.2" in paragraph.text:
                    replace_runs(paragraph, paragraph.text.replace("V1.2", "V1.3"))

    document.core_properties.title = "ArchFlow AI 产品需求文档 V1.3"
    document.core_properties.subject = "企业建筑设计 AI 提效平台面试演示 POC"
    document.core_properties.comments = "V1.3 同步内置模型预设、双 API Key 输入与公开部署边界。"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    update_document()
