from pathlib import Path
import shutil

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "ArchFlow_AI_PRD_v1.1.docx"
OUTPUT = ROOT / "docs" / "ArchFlow_AI_PRD_v1.2.docx"


def replace_runs(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_cell(cell, text):
    replace_runs(cell.paragraphs[0], text)
    for paragraph in cell.paragraphs[1:]:
        replace_runs(paragraph, "")


def update_document():
    shutil.copyfile(SOURCE, OUTPUT)
    document = Document(OUTPUT)
    paragraphs = document.paragraphs

    replacements = {
        2: "版本：V1.2  |  产品形态：企业设计师 Web 工作台  |  文档状态：7 天面试演示 POC",
        20: "首页用于展示产品定位、六个能力入口、本周项目概况与项目备忘录，重点强调快速进入任务和团队工作信息的即时记录。",
        24: "本周概况卡：展示设计进度、生成数量、资产复用与活跃趋势；支持在底部创建项目备忘录，上方最多显示两条，第 3 条起通过“更多”进入备忘录详情。",
        76: "必填：SketchUp / Rhino 模型截图、白模图或简单材质模型图；没有参考图时不允许生成",
        78: "建议：渲染风格、材质、天气、季节、景观和人物密度等简短要求",
        80: "1 张 AI 渲染效果图（MVP），不同时生成多个候选结果",
        81: "支持原图 / 生成图滑杆对比、查看生成大图、下载和重新生成",
        94: "3.8 我的资产",
        95: "定位：个人生成资产与历史记录中心，不承担企业项目管理。当前面试 POC 的资产只存在页面内存，刷新或关闭后消失。",
        96: "“我的资产”首页采用磨砂玻璃文件夹卡片，展示类型、数量与最近生成内容；支持搜索、筛选、详情、保存和二次确认删除。演示闭环为“生成 → 保存 → 在我的资产中出现 → 查看/删除”。",
        104: "六个 AI 功能统一采用“对话任务框 + 快捷条件标签 + 临时附件 + 右侧填写提示”的输入结构。设计师可以像向项目组交代任务一样描述需求，同时通过专业提示补齐关键约束。",
        105: "对话任务框：支持自然语言输入；涉及图像的模块可直接在输入框使用 Ctrl + V 粘贴图片，首屏给出与当前模块相关的专业示例。",
        109: "参考文件：支持选择文件或粘贴图片，输入区上方显示缩略图、文件数量与单项删除。文件只存在当前模块内存，离开模块或关闭页面即释放；真实生成时会发送给用户配置的第三方模型服务，ArchFlow POC 不留存。",
        110: "结果追溯：生成结果记录所属项目、功能来源、任务描述、生成时间和版本线索；保存后进入当前页面内存中的“我的资产”对应项目包。跨设备与长期留存属于后续服务端范围。",
        112: "当前目标是在约 7 天内完成可用于面试展示的 Web MVP。方案灵感、方案设计与 AI 渲染接入用户自带的真实 API；图纸美化、AI 建模与 AI 汇报保持高完成度演示数据。语言与图像 API Key 只保存在 sessionStorage，关闭标签页后清除。",
        116: "至少演示方案灵感、方案设计与 AI 渲染三条完整流程，其中渲染必须体现参考图上传与前后对比。",
        117: "“我的资产”能展示、筛选和删除当前会话生成资产，证明产品具备生成到资产沉淀的交互闭环，同时诚实说明 POC 不做长期存储。",
        118: "能清楚说明三条真实 API 链路、三条演示链路，以及对象存储、数据库、身份体系和任务队列的后续技术路线。",
    }
    for index, text in replacements.items():
        replace_runs(paragraphs[index], text)

    # Asset matrix: render assets are one generated image plus the original comparison source.
    replace_cell(document.tables[1].rows[5].cells[2], "原始参考图、单张渲染效果图、对比关系、风格、生成时间")

    # Input guidance matrix: make the render reference image explicitly required.
    replace_cell(document.tables[2].rows[5].cells[1], "说明白模、视角、材质、时间氛围与必须保持不变的主体关系；只生成一张结果。")
    replace_cell(document.tables[2].rows[5].cells[2], "必填：白模或原始效果图；建议：材质、天气、季节、景观和人物密度；输出使用滑杆与原图对比。")

    # Delivery boundaries now reflect the implemented interview POC.
    replace_cell(document.tables[3].rows[1].cells[0], "漂亮、统一的 Blueprint 磨砂玻璃 UI")
    replace_cell(document.tables[3].rows[2].cells[0], "会话级图片选择、粘贴、缩略图预览与大图查看")
    replace_cell(document.tables[3].rows[3].cells[0], "方案灵感、方案设计与 AI 渲染三条真实 API 适配")
    replace_cell(document.tables[3].rows[4].cells[0], "脱敏演示身份与用户自带 API 配置")
    replace_cell(document.tables[3].rows[5].cells[0], "生成结果保存、筛选、详情与删除闭环")
    replace_cell(document.tables[3].rows[7].cells[0], "本周概况与项目备忘录创建 / 更多详情")

    for section in document.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                if "V1.1" in paragraph.text:
                    replace_runs(paragraph, paragraph.text.replace("V1.1", "V1.2"))

    document.core_properties.title = "ArchFlow AI 产品需求文档 V1.2"
    document.core_properties.subject = "企业建筑设计 AI 提效平台面试演示 POC"
    document.core_properties.comments = "V1.2 同步真实 API、单图渲染对比、会话附件、备忘录与资产删除。"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    update_document()
